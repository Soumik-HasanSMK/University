#!/usr/bin/env python3
"""Generate Lab2_report.docx from Lab2_report.md using python-docx.

This keeps structure simple: headings map to Word headings, paragraphs map to normal text.
"""
from docx import Document
import os
from docx.shared import Pt
import re

script_dir = os.path.dirname(os.path.abspath(__file__))
INPUT = os.path.join(script_dir, '..', 'Lab2_report.md')
OUTPUT = os.path.join(script_dir, '..', 'Lab2_report.docx')


def md_to_docx(input_path=INPUT, output_path=OUTPUT):
    with open(input_path, 'r', encoding='utf-8') as f:
        md = f.read()

    doc = Document()

    # Use a simple parser: headings (#, ##) and paragraphs
    lines = md.splitlines()
    buf = []

    def flush_buf():
        nonlocal buf
        if not buf:
            return
        text = "\n".join(buf).strip()
        # Convert code blocks and inline code backticks to plain text
        text = re.sub(r'```.*?```', lambda m: m.group(0).replace('`', ''), text, flags=re.S)
        # Replace single backticks
        text = text.replace('`', '')
        for paragraph in text.split('\n\n'):
            p = doc.add_paragraph(paragraph)
            p.style.font.size = Pt(11)
        buf = []

    for line in lines:
        if line.strip() == "":
            # paragraph break
            buf.append("")
            continue

        if line.startswith('#'):
            flush_buf()
            lvl = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if lvl == 1:
                h = doc.add_heading(text, level=1)
            else:
                h = doc.add_heading(text, level=2)
            continue

        # not a heading — normal line
        buf.append(line)

    flush_buf()
    try:
        doc.save(output_path)
        print(f"Generated: {output_path}")
    except PermissionError:
        # If the target file is locked (open in another program) fall back to a new filename
        fallback = output_path.replace('.docx', '_new.docx')
        doc.save(fallback)
        print(f"Target was locked; saved to fallback: {fallback}")


if __name__ == '__main__':
    md_to_docx()
